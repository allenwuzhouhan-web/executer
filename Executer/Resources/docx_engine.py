#!/usr/bin/env python3
"""
Executer DOCX Engine
Creates .docx files from JSON spec + optional style settings.

Usage:
    python3 docx_engine.py --spec spec.json --output output.docx
    echo '{"sections":[...]}' | python3 docx_engine.py --stdin --output output.docx
"""
import argparse, importlib.util, json, os, sys
from pathlib import Path

# Import image_utils from same directory
try:
    from image_utils import resolve_image, cleanup_temp_images
except ImportError:
    _iu_path = Path(__file__).parent / "image_utils.py"
    if _iu_path.exists():
        _spec = importlib.util.spec_from_file_location("image_utils", _iu_path)
        _mod = importlib.util.module_from_spec(_spec)
        _spec.loader.exec_module(_mod)
        resolve_image = _mod.resolve_image
        cleanup_temp_images = _mod.cleanup_temp_images
    else:
        resolve_image = lambda source, temp_dir=None: source if source and Path(source).exists() else None
        cleanup_temp_images = lambda d: None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print(json.dumps({"success": False, "error": "python-docx not installed. Run: pip3 install python-docx"}))
    sys.exit(1)

def hex_to_rgb(hex_str):
    h = hex_str.lstrip("#")
    if len(h) != 6: return RGBColor(0, 0, 0)
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

def log(msg):
    print(msg, file=sys.stderr)

DEFAULTS = {
    "font": "Times New Roman",
    "heading_font": "Helvetica Neue",
    "font_size": 12,
    "heading_sizes": {0: 24, 1: 18, 2: 14, 3: 12},
    "line_spacing": 1.15,
    "margins": {"top": 1, "bottom": 1, "left": 1.25, "right": 1.25},
    "color_text": "333333",
    "color_heading": "1A1A1A",
    "color_accent": "0066CC",
}

def create_docx(spec, output_path):
    try:
        import tempfile as _tf
        _img_temp_dir = _tf.mkdtemp(prefix="executer_docx_imgs_")
        style = {**DEFAULTS, **(spec.get("style", {}))}
        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(style["margins"].get("top", 1))
            section.bottom_margin = Inches(style["margins"].get("bottom", 1))
            section.left_margin = Inches(style["margins"].get("left", 1.25))
            section.right_margin = Inches(style["margins"].get("right", 1.25))

        sections = spec.get("sections", [])
        for sec in sections:
            heading = sec.get("heading", "")
            level = sec.get("level", 1)
            body = sec.get("body", "")
            bullets = sec.get("bullets", [])
            numbered = sec.get("numbered_list", [])
            table_data = sec.get("table", None)

            # Heading
            if heading:
                h = doc.add_heading(heading, level=min(level, 4))
                for run in h.runs:
                    run.font.name = style.get("heading_font", "Helvetica Neue")
                    run.font.color.rgb = hex_to_rgb(style.get("color_heading", "1A1A1A"))

            # Body paragraph
            if body:
                p = doc.add_paragraph(body)
                p.paragraph_format.line_spacing = style.get("line_spacing", 1.15)
                for run in p.runs:
                    run.font.name = style.get("font", "Times New Roman")
                    run.font.size = Pt(style.get("font_size", 12))

            # Bullet list
            for bullet in bullets:
                p = doc.add_paragraph(bullet, style='List Bullet')
                for run in p.runs:
                    run.font.name = style.get("font", "Times New Roman")
                    run.font.size = Pt(style.get("font_size", 12))

            # Numbered list
            for item in numbered:
                p = doc.add_paragraph(item, style='List Number')
                for run in p.runs:
                    run.font.name = style.get("font", "Times New Roman")
                    run.font.size = Pt(style.get("font_size", 12))

            # Table
            if table_data:
                headers = table_data.get("headers", [])
                rows = table_data.get("rows", [])
                if headers or rows:
                    num_cols = len(headers) if headers else (len(rows[0]) if rows and rows[0] else 1)
                    total_rows = (1 if headers else 0) + len(rows)
                    table = doc.add_table(rows=total_rows, cols=num_cols)
                    table.style = 'Table Grid'

                    row_idx = 0
                    if headers:
                        for i, h in enumerate(headers):
                            cell = table.cell(0, i)
                            cell.text = str(h)
                            # Bold header
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.bold = True
                                    run.font.name = style.get("font", "Times New Roman")
                        row_idx = 1

                    for data_row in rows:
                        for i, val in enumerate(data_row):
                            if i < num_cols:
                                cell = table.cell(row_idx, i)
                                cell.text = str(val)
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.font.name = style.get("font", "Times New Roman")
                        row_idx += 1

                    doc.add_paragraph()  # spacing after table

            # Image
            image_data = sec.get("image", None)
            if image_data:
                if isinstance(image_data, str):
                    image_data = {"url": image_data}
                img_source = image_data.get("url") or image_data.get("path", "")
                local_path = resolve_image(img_source, _img_temp_dir)
                if local_path and Path(local_path).exists():
                    img_width = image_data.get("width", 5.0)
                    img_height = image_data.get("height", None)
                    kwargs = {"image_path_or_stream": local_path}
                    if img_width:
                        kwargs["width"] = Inches(img_width)
                    if img_height:
                        kwargs["height"] = Inches(img_height)
                    doc.add_picture(**kwargs)
                    # Set alignment
                    alignment_str = image_data.get("alignment", "center").lower()
                    align_map = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = align_map.get(alignment_str, WD_ALIGN_PARAGRAPH.CENTER)
                    # Caption
                    caption = image_data.get("caption", "")
                    if caption:
                        cap_para = doc.add_paragraph(caption)
                        cap_para.alignment = align_map.get(alignment_str, WD_ALIGN_PARAGRAPH.CENTER)
                        for run in cap_para.runs:
                            run.font.italic = True
                            run.font.size = Pt(style.get("font_size", 12) - 2)
                            run.font.name = style.get("font", "Times New Roman")
                            run.font.color.rgb = hex_to_rgb(style.get("color_text", "333333"))

        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out))
        cleanup_temp_images(_img_temp_dir)
        log(f"Saved: {out}")
        return {"success": True, "path": str(out.resolve()), "sections": len(sections)}
    except Exception as e:
        return {"success": False, "error": f"{type(e).__name__}: {str(e)}"}

def main():
    parser = argparse.ArgumentParser(description="Executer DOCX Engine")
    parser.add_argument("--spec", help="Path to spec JSON file")
    parser.add_argument("--stdin", action="store_true", help="Read spec from stdin")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    if args.stdin:
        try: spec = json.load(sys.stdin)
        except json.JSONDecodeError as e:
            print(json.dumps({"success": False, "error": f"Invalid JSON: {e}"})); sys.exit(1)
    elif args.spec:
        try: spec = json.loads(Path(args.spec).read_text(encoding="utf-8"))
        except Exception as e:
            print(json.dumps({"success": False, "error": str(e)})); sys.exit(1)
    else:
        print(json.dumps({"success": False, "error": "Provide --spec or --stdin"})); sys.exit(1)

    output = args.output
    if output.endswith("/") or Path(output).is_dir():
        filename = spec.get("filename", "document.docx")
        if not filename.endswith(".docx"): filename += ".docx"
        output = str(Path(output) / filename)

    result = create_docx(spec, output)
    print(json.dumps(result))
    sys.exit(0 if result["success"] else 1)

if __name__ == "__main__":
    main()
